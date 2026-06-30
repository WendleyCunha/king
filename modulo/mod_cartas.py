import streamlit as st
import pandas as pd
import zipfile
import os
import time
from datetime import datetime
from docx import Document
from io import BytesIO

from database import (
    obter_base_colaboradores_db, salvar_novo_colaborador_db, deletar_colaborador_db,
    obter_cartas_db, criar_carta_db, atualizar_carta_db,
    registrar_assinatura_carta_db, deletar_carta_db, reabrir_carta_db,
    fechar_lote_cartas_db, listar_lotes_cartas_db, limpar_anexos_lote_db,
    tem_permissao, pode_exportar, pode_deletar,
)


# ─── GERAÇÃO DE WORD ──────────────────────────────────────────────────────────

def gerar_word_memoria(dados):
    try:
        diretorio_raiz = os.getcwd()
        template_path = os.path.join(diretorio_raiz, "carta_preenchida.docx")
        if not os.path.exists(template_path):
            template_path = "carta_preenchida.docx"

        doc = Document(template_path)

        for p in doc.paragraphs:
            for k, v in dados.items():
                if f"{{{{{k}}}}}" in p.text:
                    p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in dados.items():
                            if f"{{{{{k}}}}}" in p.text:
                                p.text = p.text.replace(f"{{{{{k}}}}}", str(v))

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"⚠️ Erro no Template Word: {e}")
        return None


def _dados_word(c):
    return {
        "GERENTE":        c["NOME"],
        "CPF":            c["CPF"],
        "CODIGO_CLIENTE": c.get("COD_CLI", ""),
        "VALOR_DEBITO":   f"R$ {c['VALOR']:,.2f}",
        "LOJA_ORIGEM":    c.get("LOJA", ""),
        "DATA_COMPRA":    c.get("DATA", ""),
        "DESC_DEBITO":    c.get("MOTIVO", ""),
        "DATA_LOCAL":     f"São Paulo, {datetime.now().strftime('%d/%m/%Y')}",
    }


# ─── EDIÇÃO DE CARTA ──────────────────────────────────────────────────────────

def _renderizar_form_edicao(c, dict_colab):
    """Renderiza o formulário inline de edição de uma carta.
    Salva via atualizar_carta_db e fecha o painel ao confirmar."""
    carta_id = c["id"]

    with st.form(key=f"form_edit_{carta_id}", clear_on_submit=False):
        st.markdown("**✏️ Editar Carta**")

        lista_nomes = sorted(dict_colab.keys())
        opcoes = ["+ CADASTRAR NOVO"] + lista_nomes
        idx_atual = opcoes.index(c["NOME"]) if c["NOME"] in opcoes else 0
        escolha = st.selectbox("Colaborador", opcoes, index=idx_atual,
                               key=f"ed_nome_sel_{carta_id}")

        e1, e2, e3 = st.columns(3)

        if escolha == "+ CADASTRAR NOVO":
            novo_nome = e1.text_input("Nome *", key=f"ed_nome_{carta_id}").upper().strip()
            novo_cpf  = e2.text_input("CPF *",  key=f"ed_cpf_{carta_id}",
                                      value=c.get("CPF", ""))
        else:
            novo_nome = escolha
            novo_cpf  = e2.text_input("CPF", key=f"ed_cpf_{carta_id}",
                                      value=dict_colab.get(escolha, c.get("CPF", "")))

        novo_cod = e3.text_input("Código do Cliente", key=f"ed_cod_{carta_id}",
                                 value=c.get("COD_CLI", ""))

        f1, f2, f3 = st.columns(3)
        novo_valor = f1.number_input("Valor R$", min_value=0.0, step=0.01,
                                     value=float(c.get("VALOR", 0)),
                                     key=f"ed_valor_{carta_id}")
        nova_loja  = f2.text_input("Loja Origem", key=f"ed_loja_{carta_id}",
                                   value=c.get("LOJA", "")).upper()

        # Converte data string "dd/mm/yyyy" → objeto date para o widget
        data_str = c.get("DATA", "")
        try:
            data_default = datetime.strptime(data_str, "%d/%m/%Y").date()
        except Exception:
            data_default = datetime.today().date()
        nova_data = f3.date_input("Data da Ocorrência", value=data_default,
                                  key=f"ed_data_{carta_id}")

        novo_motivo = st.text_area("Motivo Detalhado", key=f"ed_motivo_{carta_id}",
                                   value=c.get("MOTIVO", "")).upper()

        col_salvar, col_cancelar = st.columns(2)
        salvar    = col_salvar.form_submit_button("💾 Salvar Alterações",
                                                  type="primary",
                                                  use_container_width=True)
        cancelar  = col_cancelar.form_submit_button("✖ Cancelar",
                                                    use_container_width=True)

        if salvar:
            nome_final = novo_nome if escolha == "+ CADASTRAR NOVO" else escolha
            if not nome_final or not novo_cpf or not novo_cod:
                st.error("Preencha Nome, CPF e Código do Cliente.")
            else:
                if escolha == "+ CADASTRAR NOVO":
                    salvar_novo_colaborador_db(nome_final, novo_cpf)
        
                atualizar_carta_db(          # ← campos em MAIÚSCULAS = nomes reais no Firestore
                    carta_id,
                    NOME=nome_final,
                    CPF=novo_cpf,
                    COD_CLI=novo_cod,
                    VALOR=novo_valor,
                    LOJA=nova_loja,
                    DATA=nova_data.strftime("%d/%m/%Y"),
                    MOTIVO=novo_motivo,
                )
                st.session_state.pop(f"editando_{carta_id}", None)
                st.success("✅ Carta atualizada!")
                st.rerun()


# ─── GERAÇÃO DE ZIP (arquivos ASSINADOS realmente enviados no upload) ────────

def gerar_zip_lote(cartas_lote):
    """Recebe uma lista de cartas e retorna (bytes_do_zip, quantidade_incluida)
    com os ARQUIVOS ASSINADOS de verdade (o que foi enviado no upload de cada
    carta, guardado em anexo_bin/nome_arquivo) — NÃO regenera o template em
    branco. Cartas sem arquivo enviado (ainda não assinadas, ou já tiveram o
    anexo limpo do banco) são ignoradas."""
    zip_buffer = BytesIO()
    incluidas = 0
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for c in cartas_lote:
            anexo = c.get("anexo_bin")
            if not anexo:
                continue  # ainda não tem assinatura enviada (ou já foi limpa)
            nome_original = c.get("nome_arquivo") or f"{c.get('NOME','arquivo')}.docx"
            ext = os.path.splitext(nome_original)[1] or ".docx"
            nome_arquivo = f"Carta_{c.get('NOME','').replace(' ', '_')}_{c.get('id','')}{ext}"
            zf.writestr(nome_arquivo, anexo)
            incluidas += 1
    zip_buffer.seek(0)
    return zip_buffer.getvalue(), incluidas


# ─── GERAÇÃO DE EXCEL ────────────────────────────────────────────────────────

def gerar_excel_lote(cartas_lote):
    colunas = ["NOME", "CPF", "COD_CLI", "VALOR", "LOJA", "DATA", "MOTIVO", "status", "id_lote"]
    rows = [{col: c.get(col, "") for col in colunas} for c in cartas_lote]

    df = pd.DataFrame(rows, columns=colunas)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Cartas")
        ws = writer.sheets["Cartas"]
        for i, col in enumerate(df.columns):
            largura = max(df[col].astype(str).map(len).max(), len(col)) + 2
            ws.set_column(i, i, largura)
    buf.seek(0)
    return buf.getvalue()


# ─── ESTILO (segue a paleta/identidade do main.py — dourado #C9A84C) ──────────

_CSS_CARTAS = """
<style>
.cartas-card {
    background:#fff; border:1px solid #e2e8f0; border-radius:12px;
    padding:16px; margin-bottom:6px; border-top:4px solid #C9A84C;
    box-shadow:0 2px 8px rgba(0,0,0,0.04); min-height:170px;
}
.cartas-loja-header {
    background:#2c3e50; color:#C9A84C; padding:8px 15px;
    border-radius:8px; margin:22px 0 14px 0; font-weight:700;
}
.cartas-label { color:#64778d; font-size:0.78rem; font-weight:700;
    text-transform:uppercase; letter-spacing:0.5px; margin-bottom:2px; }
.cartas-info  { font-weight:700; color:#2c3e50; margin-bottom:8px; }
.cartas-edit-box {
    background:#f8f9fa; border:1px solid #C9A84C; border-radius:10px;
    padding:16px; margin-top:6px; margin-bottom:10px;
}
</style>
"""


# ─── INTERFACE PRINCIPAL ──────────────────────────────────────────────────────

def renderizar_cartas(papel, user=None):
    st.markdown(_CSS_CARTAS, unsafe_allow_html=True)
    st.subheader("📑 Gestão de Cartas de Débito")

    pode_exp = pode_exportar(user) if user else (papel in ("adm", "supervisor"))
    pode_del = pode_deletar(user) if user else (papel == "adm")

    cartas      = obter_cartas_db()
    dict_colab  = obter_base_colaboradores_db()
    lista_nomes = sorted(dict_colab.keys())

    tabs = st.tabs(["🆕 Nova Carta", "📋 Painel", "📦 Fechamento", "✅ Histórico", "⚙️ Config"])

    # ── ABA 0: NOVA CARTA ────────────────────────────────────────────────────
    with tabs[0]:
        st.markdown("##### Informações do Lançamento")
        escolha_nome = st.selectbox(
            "Busque ou selecione o Colaborador:", ["+ CADASTRAR NOVO"] + lista_nomes
        )

        with st.form("f_nova_carta", clear_on_submit=True):
            c1, c2, c3 = st.columns(3)
            if escolha_nome == "+ CADASTRAR NOVO":
                nome = c1.text_input("Nome do Colaborador").upper().strip()
                cpf  = c2.text_input("CPF")
            else:
                nome = escolha_nome
                cpf  = c2.text_input("CPF", value=dict_colab.get(escolha_nome, ""))

            cod_cli = c3.text_input("Código do Cliente")
            v1, v2, v3 = st.columns(3)
            valor  = v1.number_input("Valor R$", min_value=0.0, step=0.01)
            loja   = v2.text_input("Loja Origem").upper()
            data_c = v3.date_input("Data da Ocorrência")
            motivo = st.text_area("Motivo Detalhado").upper()

            if st.form_submit_button("✨ Gerar e Registrar", type="primary", use_container_width=True):
                if nome and cpf and cod_cli:
                    if escolha_nome == "+ CADASTRAR NOVO":
                        salvar_novo_colaborador_db(nome, cpf)

                    criar_carta_db(
                        nome=nome, cpf=cpf, cod_cli=cod_cli, valor=valor,
                        loja=loja, data_str=data_c.strftime("%d/%m/%Y"), motivo=motivo,
                    )
                    st.success("✅ Carta registrada com sucesso!")
                    st.rerun()
                else:
                    st.error("Preencha Nome, CPF e Código do Cliente.")

    # ── ABA 1: PAINEL ────────────────────────────────────────────────────────
    with tabs[1]:
        lista_painel = [c for c in cartas if c.get("status") == "Aguardando Assinatura"]

        if lista_painel:
            total_valor = sum(c.get("VALOR", 0) for c in lista_painel)
            m1, m2 = st.columns(2)
            m1.markdown(
                f'<div class="kpi-card gold"><div class="kpi-label">Pendentes</div>'
                f'<div class="kpi-value">{len(lista_painel)}</div></div>',
                unsafe_allow_html=True,
            )
            m2.markdown(
                f'<div class="kpi-card gold"><div class="kpi-label">Valor Total</div>'
                f'<div class="kpi-value">R$ {total_valor:,.2f}</div></div>',
                unsafe_allow_html=True,
            )
            st.write("")

            busca_p = st.text_input("🔍 Buscar (Nome / Código)")
            if busca_p:
                lista_painel = [
                    c for c in lista_painel
                    if busca_p.upper() in c["NOME"] or busca_p in str(c.get("COD_CLI", ""))
                ]

            df_p = pd.DataFrame(lista_painel).sort_values(by="LOJA")

            for loja_n, group in df_p.groupby("LOJA"):
                st.markdown(
                    f'<div class="cartas-loja-header">📍 {loja_n} ({len(group)} itens)</div>',
                    unsafe_allow_html=True,
                )
                cols = st.columns(3)
                for idx, (_, c) in enumerate(group.iterrows()):
                    carta_id   = c["id"]
                    editando   = st.session_state.get(f"editando_{carta_id}", False)

                    with cols[idx % 3]:
                        # ── Card visual ──────────────────────────────────────
                        st.markdown(
                            f'<div class="cartas-card">'
                            f'<div class="cartas-label">Colaborador</div>'
                            f'<div class="cartas-info">{c["NOME"]}</div>'
                            f'<div class="cartas-label">Cód. Cliente</div>'
                            f'<div class="cartas-info">{c.get("COD_CLI","—")}</div>'
                            f'<div class="cartas-label">Valor</div>'
                            f'<div style="font-size:1.2rem;font-weight:800;color:#C9A84C;margin-bottom:6px;">'
                            f'R$ {c["VALOR"]:,.2f}</div>'
                            f'<div class="cartas-label">Data: {c["DATA"]}</div>'
                            f'<div class="cartas-label" style="margin-top:8px;">Motivo</div>'
                            f'<div style="font-size:0.85rem;color:#495057;">{c.get("MOTIVO","—")}</div>'
                            f"</div>",
                            unsafe_allow_html=True,
                        )

                        # ── Botões de ação (Baixar | Editar | Deletar) ───────
                        w_bytes = gerar_word_memoria(_dados_word(c))

                        if pode_del:
                            btn_baixar, btn_editar, btn_del = st.columns(3)
                        else:
                            btn_baixar, btn_editar = st.columns(2)
                            btn_del = None

                        if w_bytes:
                            btn_baixar.download_button(
                                "📂 Baixar",
                                w_bytes,
                                file_name=f"Carta_{c['NOME']}.docx",
                                key=f"w_{carta_id}",
                                use_container_width=True,
                            )

                        # Botão Editar: alterna visibilidade do formulário
                        label_edit = "✖ Fechar" if editando else "✏️ Editar"
                        if btn_editar.button(label_edit, key=f"btn_edit_{carta_id}",
                                             use_container_width=True):
                            st.session_state[f"editando_{carta_id}"] = not editando
                            st.rerun()

                        if btn_del and btn_del.button("🗑️", key=f"del_{carta_id}",
                                                      use_container_width=True):
                            deletar_carta_db(carta_id)
                            st.rerun()

                        # ── Formulário de edição inline ──────────────────────
                        if st.session_state.get(f"editando_{carta_id}", False):
                            with st.container():
                                st.markdown('<div class="cartas-edit-box">',
                                            unsafe_allow_html=True)
                                _renderizar_form_edicao(c, dict_colab)
                                st.markdown("</div>", unsafe_allow_html=True)

                        # ── Upload da carta assinada ─────────────────────────
                        up = st.file_uploader(
                            "Upload Assinada", key=f"up_{carta_id}",
                            label_visibility="collapsed"
                        )
                        if up:
                            registrar_assinatura_carta_db(carta_id, up.getvalue(), up.name)
                            st.success("✅ Recebida!")
                            st.rerun()
        else:
            st.info("Nenhuma carta aguardando assinatura.")

    # ── ABA 2: FECHAMENTO DE LOTE ────────────────────────────────────────────
    with tabs[2]:
        prontas = [c for c in cartas if c.get("status") == "CARTA RECEBIDA"]

        if not prontas:
            st.info("Nenhuma carta assinada pronta para fechamento.")
        else:
            st.markdown(f"##### 📦 Lote pronto: {len(prontas)} itens")

            df_prontas = pd.DataFrame(prontas)[["NOME", "VALOR", "LOJA", "COD_CLI", "DATA"]]
            st.dataframe(df_prontas, use_container_width=True)

            total_lote = sum(c.get("VALOR", 0) for c in prontas)
            st.markdown(
                f'<div class="kpi-card gold" style="max-width:320px;">'
                f'<div class="kpi-label">Valor Total do Lote</div>'
                f'<div class="kpi-value">R$ {total_lote:,.2f}</div></div>',
                unsafe_allow_html=True,
            )
            st.divider()

            if pode_exp:
                col_zip, col_xls = st.columns(2)

                zip_bytes, n_assinadas = gerar_zip_lote(prontas)
                col_zip.download_button(
                    f"📥 Baixar ZIP (arquivos assinados — {n_assinadas}/{len(prontas)})",
                    data=zip_bytes,
                    file_name=f"Lote_{datetime.now().strftime('%Y%m%d_%H%M')}_Assinadas.zip",
                    mime="application/zip",
                    use_container_width=True,
                    disabled=(n_assinadas == 0),
                )
                if n_assinadas < len(prontas):
                    col_zip.caption("⚠️ Algumas cartas ainda não têm o arquivo assinado anexado.")

                excel_bytes = gerar_excel_lote(prontas)
                col_xls.download_button(
                    "📊 Baixar Excel do Lote",
                    data=excel_bytes,
                    file_name=f"Lote_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
                st.divider()

            if st.button(
                "🚀 FINALIZAR LOTE E ENVIAR AO HISTÓRICO",
                type="primary",
                use_container_width=True,
            ):
                id_lote = fechar_lote_cartas_db(prontas)
                st.success(f"✅ Lote {id_lote} finalizado! {len(prontas)} cartas arquivadas.")
                st.rerun()

    # ── ABA 3: HISTÓRICO ────────────────────────────────────────────────────
    with tabs[3]:
        st.markdown("##### ✅ Histórico de Lotes Fechados")

        lotes = listar_lotes_cartas_db()

        if not lotes:
            st.info("Nenhum lote fechado ainda.")
        else:
            for lote in lotes:
                with st.expander(
                    f"📦 Lote {lote['id']}  ·  {lote.get('data','—')}  ·  "
                    f"{lote.get('total', 0)} cartas  ·  "
                    f"R$ {lote.get('valor_total', 0):,.2f}"
                ):
                    ids_lote = lote.get("ids_cartas", [])
                    cartas_lote = [c for c in cartas if c.get("id") in ids_lote]

                    if not cartas_lote:
                        st.caption("Cartas deste lote não encontradas no banco.")
                        continue

                    st.dataframe(
                        pd.DataFrame(cartas_lote)[["NOME", "VALOR", "LOJA", "COD_CLI", "DATA"]],
                        use_container_width=True,
                    )

                    n_com_anexo = sum(1 for c in cartas_lote if c.get("anexo_bin"))
                    col_x, col_z, col_del = st.columns(3)

                    # ── Botão 1: Excel completo ──
                    if pode_exp:
                        xls_hist = gerar_excel_lote(cartas_lote)
                        col_x.download_button(
                            "📊 Excel Completo",
                            data=xls_hist,
                            file_name=f"Lote_{lote['id']}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"xls_hist_{lote['id']}",
                            use_container_width=True,
                        )
                    else:
                        col_x.caption("🔒 Exportação restrita.")

                    # ── Botão 2: ZIP com os arquivos ASSINADOS ──
                    if pode_exp:
                        zip_hist, _ = gerar_zip_lote(cartas_lote)
                        col_z.download_button(
                            f"📥 ZIP Assinadas ({n_com_anexo}/{len(cartas_lote)})",
                            data=zip_hist,
                            file_name=f"Lote_{lote['id']}_Assinadas.zip",
                            mime="application/zip",
                            key=f"zip_hist_{lote['id']}",
                            use_container_width=True,
                            disabled=(n_com_anexo == 0),
                        )
                    else:
                        col_z.caption("🔒 Exportação restrita.")

                    # ── Botão 3: Excluir os anexos do lote ──
                    with col_del:
                        if pode_del:
                            if n_com_anexo == 0:
                                st.caption("✅ Este lote já não tem arquivos anexados.")
                            else:
                                conf_key = f"confdel_{lote['id']}"
                                confirmar = st.checkbox(
                                    f"Confirmo excluir os {n_com_anexo} arquivo(s)",
                                    key=conf_key,
                                )
                                if st.button(
                                    "🗑️ Excluir Arquivos do Lote",
                                    key=f"delanexos_{lote['id']}",
                                    use_container_width=True,
                                    disabled=not confirmar,
                                ):
                                    limpar_anexos_lote_db(ids_lote)
                                    st.success(
                                        "✅ Arquivos anexados removidos do banco! "
                                        "O histórico (nomes, valores, status, datas) foi mantido."
                                    )
                                    time.sleep(1)
                                    st.rerun()
                                st.caption(
                                    "⚠️ Baixe o ZIP acima antes de excluir — esta ação é "
                                    "irreversível e só remove o arquivo, não o registro do lote."
                                )
                        else:
                            st.caption("🔒 Exclusão restrita a administradores.")

    # ── ABA 4: CONFIG ────────────────────────────────────────────────────────
    with tabs[4]:
        if papel not in ("adm", "supervisor"):
            st.warning("🔒 Acesso restrito a Administradores e Supervisores.")
            return

        st.markdown("##### ⚙️ Configurações")

        with st.expander("👤 Gerenciar Colaboradores", expanded=True):
            st.caption(f"{len(dict_colab)} colaborador(es) cadastrado(s)")

            with st.form("form_colab_cartas", clear_on_submit=True):
                cc1, cc2 = st.columns(2)
                novo_nome_c = cc1.text_input("Nome *")
                novo_cpf_c  = cc2.text_input("CPF *")
                if st.form_submit_button("➕ Adicionar", use_container_width=True):
                    if novo_nome_c and novo_cpf_c:
                        salvar_novo_colaborador_db(novo_nome_c, novo_cpf_c)
                        st.success(f"✅ {novo_nome_c.upper()} adicionado!")
                        st.rerun()
                    else:
                        st.error("Preencha nome e CPF.")

            st.markdown("---")
            st.markdown("**📤 Importar em lote (Excel ou CSV)**")
            st.caption("Coluna A = Nome, Coluna B = CPF. Aceita planilha com ou sem cabeçalho.")
            up_colab = st.file_uploader(
                "Selecione o arquivo", type=["xlsx", "csv"], key="up_colab_lote"
            )
            if up_colab is not None:
                try:
                    if up_colab.name.lower().endswith(".csv"):
                        df_up = pd.read_csv(up_colab, header=None, dtype=str)
                    else:
                        df_up = pd.read_excel(up_colab, header=None, dtype=str)

                    if df_up.shape[1] < 2:
                        st.error("O arquivo precisa ter pelo menos 2 colunas (Nome e CPF).")
                    else:
                        df_up = df_up.iloc[:, :2]
                        df_up.columns = ["NOME", "CPF"]
                        df_up = df_up.dropna(how="all")

                        st.markdown(f"**Prévia ({len(df_up)} linha(s) encontradas):**")
                        st.dataframe(df_up, use_container_width=True, hide_index=True)

                        tem_cabecalho = st.checkbox(
                            "A primeira linha é um cabeçalho (ex: 'Nome', 'CPF') — ignorar",
                            value=True, key="up_colab_header",
                        )
                        if st.button("✅ Confirmar importação", key="btn_import_colab",
                                     type="primary", use_container_width=True):
                            linhas = df_up.iloc[1:] if tem_cabecalho else df_up
                            importados, ignorados = 0, 0
                            for _, row in linhas.iterrows():
                                nome_imp = str(row["NOME"]).strip() if pd.notna(row["NOME"]) else ""
                                cpf_imp  = str(row["CPF"]).strip() if pd.notna(row["CPF"]) else ""
                                if nome_imp and nome_imp.lower() != "nan" \
                                        and cpf_imp and cpf_imp.lower() != "nan":
                                    salvar_novo_colaborador_db(nome_imp, cpf_imp)
                                    importados += 1
                                else:
                                    ignorados += 1
                            msg = f"✅ {importados} colaborador(es) importado(s)!"
                            if ignorados:
                                msg += f" ({ignorados} linha(s) ignorada(s) por falta de nome/CPF)"
                            st.success(msg)
                            time.sleep(1.2)
                            st.rerun()
                except Exception as e:
                    st.error(f"⚠️ Erro ao ler o arquivo: {e}")

            st.markdown("---")
            if dict_colab:
                st.write("**Colaboradores cadastrados:**")
                for nome_c, cpf_c in sorted(dict_colab.items()):
                    cc1, cc2 = st.columns([4, 1])
                    cc1.write(f"**{nome_c}** — CPF: {cpf_c}")
                    if cc2.button("🗑️", key=f"del_colab_{nome_c}", use_container_width=True):
                        deletar_colaborador_db(nome_c)
                        st.toast(f"{nome_c} removido.")
                        st.rerun()

        with st.expander("📤 Exportar Base Completa"):
            if cartas:
                xls_all = gerar_excel_lote(cartas)
                st.download_button(
                    "📊 Exportar TODAS as cartas (.xlsx)",
                    data=xls_all,
                    file_name=f"Base_Completa_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            else:
                st.info("Nenhuma carta no banco ainda.")

        if papel == "adm":
            with st.expander("🚨 Zona de Perigo"):
                st.warning("Reabrir uma carta muda seu status de volta para 'Aguardando Assinatura'.")
                cartas_fechadas = [c for c in cartas if c.get("status") == "LOTE_FECHADO"]
                if cartas_fechadas:
                    nomes_f = [f"{c['NOME']} — {c['DATA']}" for c in cartas_fechadas]
                    idx_sel = st.selectbox("Selecionar carta para reabrir:", range(len(nomes_f)),
                                           format_func=lambda i: nomes_f[i])
                    if st.button("🔓 Reabrir Carta Selecionada", type="primary"):
                        reabrir_carta_db(cartas_fechadas[idx_sel]["id"])
                        st.success("Carta reaberta.")
                        st.rerun()
                else:
                    st.caption("Nenhuma carta fechada para reabrir.")
